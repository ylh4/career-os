#!/usr/bin/env python3
"""Render a staged Career OS markdown artifact (resume / cover letter) to .docx.

The pipeline tracks markdown sources (resume_v1.md, cover_v1.md, provenance.md);
this renders the human-facing .docx deliverable that gets uploaded to an ATS. Per
.gitignore, the .docx is a throwaway export — regenerate it any time from the .md.

Output style — "Modern sans + rule lines", ATS-safe (single column, standard font, real text,
no tables/text-boxes): Calibri base, ~0.8in margins; the first `#` is the candidate name
(large bold) with the line under it styled as a small grey contact line; `##` are bold
small-caps section headings with a thin bottom rule; `###` are smaller bold sub-headings.

Supported markdown subset (enough for resumes/letters):
  #                 -> name (first one) / large heading
  ##, ###           -> section heading (with rule) / sub-heading
  - item / * item   -> bullet list
  **bold**          -> bold runs (inline, anywhere in a line)
  blank line        -> paragraph break
  everything else   -> a normal paragraph

Requires python-docx (installed in the project venv).

Usage:
    python scripts/render_docx.py artifacts/<id>/resume_v1.md            # -> resume_v1.docx
    python scripts/render_docx.py artifacts/<id>/resume_v1.md out.docx   # explicit output
    python scripts/render_docx.py artifacts/<id>/*.md                    # batch
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")

GREY = RGBColor(0x66, 0x66, 0x66)


def add_inline(paragraph, text: str) -> None:
    """Add text to a paragraph, turning **bold** spans into bold runs."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _setup(doc: "Document") -> None:
    """Clean ATS-safe document defaults: Calibri, tight margins, compact spacing."""
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(0.8)
        section.top_margin = section.bottom_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.08


def _bottom_rule(paragraph, color: str = "999999") -> None:
    """Draw a thin bottom border under a heading paragraph (the section rule)."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # 6/8 pt ≈ 0.75pt
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def _name(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.paragraph_format.space_after = Pt(2)


def _contact(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(8)


def _section(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.small_caps = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    _bottom_rule(p)


def _subheading(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def render(md_path: Path, out_path: Path) -> None:
    doc = Document()
    _setup(doc)
    seen_name = False
    expect_contact = False  # the line right after the name is the contact line
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue  # blank lines collapse; spacing is handled by paragraph styles
        h = HEADING_RE.match(line)
        if h:
            level, text = len(h.group(1)), h.group(2).strip()
            if level == 1 and not seen_name:
                _name(doc, text)
                seen_name, expect_contact = True, True
            elif level <= 2:
                _section(doc, text)
                expect_contact = False
            else:
                _subheading(doc, text)
                expect_contact = False
            continue
        b = BULLET_RE.match(line)
        if b:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, b.group(1).strip())
            expect_contact = False
            continue
        if expect_contact:
            _contact(doc, line)
            expect_contact = False
            continue
        add_inline(doc.add_paragraph(), line)
    doc.save(str(out_path))


def main() -> int:
    args = sys.argv[1:]
    if not args:
        print(__doc__)
        return 2

    # Two-arg form: explicit input + output. Otherwise treat every arg as an input.
    pairs: list[tuple[Path, Path]] = []
    if len(args) == 2 and args[1].lower().endswith(".docx"):
        pairs.append((Path(args[0]), Path(args[1])))
    else:
        for a in args:
            src = Path(a)
            pairs.append((src, src.with_suffix(".docx")))

    rc = 0
    for src, out in pairs:
        if not src.is_file():
            print(f"error: not a file: {src}", file=sys.stderr)
            rc = 1
            continue
        render(src, out)
        print(f"rendered {src} -> {out}")
    return rc


if __name__ == "__main__":
    sys.exit(main())
